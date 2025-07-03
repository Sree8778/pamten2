import io
import re # Import the regular expression module
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from weasyprint import HTML
from jinja2 import Environment, FileSystemLoader
import os
from bs4 import BeautifulSoup

# --- NEW: Helper function to clean up extra whitespace ---
def clean_text(text: str) -> str:
    """
    Removes extra newlines and leading/trailing whitespace from a string.
    This prevents large blank gaps in the final PDF.
    """
    if not text:
        return ""
    # Replace two or more newlines with a single one
    cleaned_text = re.sub(r'\n\s*\n', '\n', text)
    return cleaned_text.strip()

# Helper to strip HTML tags for clean text in DOCX
def strip_html(html_string):
    if not html_string:
        return ""
    soup = BeautifulSoup(html_string, "html.parser")
    for br in soup.find_all("br"):
        br.replace_with("\n")
    return soup.get_text()

# --- DOCX GENERATION ---
def generate_docx_from_data(data):
    doc = Document()
    style = data.get('styleOptions', {})
    font_name = style.get('fontFamily', 'Calibri').split(',')[0]
    font_size = style.get('fontSize', 11)
    accent_color_hex = style.get('accentColor', '#34495e').lstrip('#')
    accent_color_rgb = RGBColor.from_string(accent_color_hex)

    normal_style = doc.styles['Normal']
    normal_style.font.name = font_name
    normal_style.font.size = Pt(font_size)

    heading_style = doc.styles.add_style('SectionHeading', 1)
    heading_style.font.name = font_name
    heading_style.font.size = Pt(14)
    heading_style.font.bold = True
    heading_style.font.color.rgb = accent_color_rgb
    heading_style.paragraph_format.space_before = Pt(12)
    heading_style.paragraph_format.space_after = Pt(6)

    # --- Header ---
    personal = data.get('personal', {})
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    runner = p.add_run(personal.get('name', ''))
    runner.bold = True
    runner.font.name = font_name
    runner.font.size = Pt(24)
    runner.font.color.rgb = accent_color_rgb

    contact_items = [
        personal.get('email'),
        personal.get('phone'),
        personal.get('location')
    ]
    if personal.get('legalStatus') and personal.get('legalStatus') != 'Prefer not to say':
        contact_items.append(personal.get('legalStatus'))

    contact_info = " | ".join(filter(None, contact_items))
    p = doc.add_paragraph(contact_info)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph() 

    # --- Sections (with text cleaning for DOCX as well) ---
    if data.get('summary'):
        doc.add_paragraph('Summary', style='SectionHeading')
        doc.add_paragraph(clean_text(strip_html(data.get('summary'))), style=normal_style)

    if data.get('experience') and any(e.get('jobTitle') for e in data.get('experience')):
        doc.add_paragraph('Experience', style='SectionHeading')
        for exp in data['experience']:
            p = doc.add_paragraph()
            p.add_run(exp.get('jobTitle', '')).bold = True
            p.add_run(f"\n{exp.get('company', '')} | {exp.get('dates', '')}\n").italic = True
            p.add_run(clean_text(strip_html(exp.get('description', ''))))
            p.style = normal_style
            p.paragraph_format.space_after = Pt(12)
            
    return doc


# --- PDF GENERATION ---
def generate_pdf_from_data(data):
    # --- NEW: Clean the data before rendering ---
    # We iterate through the data and apply the clean_text function to relevant fields.
    if data.get('summary'):
        data['summary'] = clean_text(data['summary'])
    
    if data.get('experience'):
        for exp in data['experience']:
            exp['description'] = clean_text(exp['description'])
            
    if data.get('education'):
        for edu in data['education']:
            edu['achievements'] = clean_text(edu['achievements'])

    # Now, we render the template with the cleaned data
    env = Environment(loader=FileSystemLoader(os.path.join(os.path.dirname(__file__), 'assets')))
    template = env.get_template('resume_template.html')
    
    rendered_html = template.render(**data)
    
    return HTML(string=rendered_html).write_pdf()